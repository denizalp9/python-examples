
# -- Gerekli Modüller -- #

import docx # Microsoft Word dosyaları üzerinde kontrol sağlamak için kullanılan modül
import os
import datetime # kısaca tarih modülü, belirtilen formatta tarih veriyor.
import locale # datetime modülü için konum seçmeye yarıyor. Örneğin: "2020 May Friday" yerine "2020 Mayıs Cuma" şeklinde türkçe tarih bilgisi alabiliyorsun.

from docx.shared import * # docx için ekstra importlar
from docx.enum.text import *

# -- Tarih Alma -- #

an = datetime.datetime.now() # Programın çalıştırıldığı güncel tarihi alır.

locale.setlocale(locale.LC_ALL, 'turkish') # Tarih bilgisini türkçeye çevirir
localtime = datetime.datetime.strftime(an, '%B %Y') # Yıl ve Ay formatında türkçe tarihi bir değişkene atadım


document = docx.Document()

# -- Sayfa Ayarları -- #

section = document.sections[0] # Word için sayfa ayarları(boyutlandırma)
section.left_margin = Inches(0.55)
section.right_margin = Inches(0.75)

# -- Tarihi Başlık Olarak Ekleme --#

hdrTime = document.add_heading(f'{localtime} bilmemne apartmanı hesap kitap bişeyler') # Header(Başlık), içindeki değişken yukarıda atadığım, program çalıştığında güncel zamanı türkçe olarak veren değişken.


# -- Tablo İşlemleri -- #


table = document.add_table(rows=1, cols=3) # Tablo oluşturma

table.style= 'Table Grid' # Oluşturduğum tablonun kenarlıkları olması için bu stili seçtim


insanlar = [
    [1, "Mehmet Aydın", "Daire: 33"],
    [2, "Ahmet Sonuç", "Daire: 3"],
    [3, "Örnek", "Daire: 1"]
]


headerCells = table.rows[0].cells # Tablonun Header(Başlık)ları için tablonun ilk satırının hücrelerini "headerCells" değişkenine atıyorum
# Her hücrenin başlığını elle giriyorum
headerCells[0].text = 'ID'
headerCells[1].text = 'İsim-Soyisim'
headerCells[2].text = 'Daire'


for ID, name, doornum in insanlar: # "insanlar" listesinin içindeki her bir kişi için yeni bir satır açıyor
    rowCells = table.add_row().cells
    rowCells[0].text = str(ID)
    rowCells[1].text = name
    rowCells[2].text = doornum


# -- Diğer İşlemler -- #

style = document.styles['Normal'] # Belgenin stilini ayarlar(Burada normali kullanıyorum)
font = style.font # Belgenin fontunu değiştirmek için "font" değişkenine atıyorum

font.name = 'Calibri'
font.size = Pt(15)


boldparg = document.add_paragraph() # Kalın yazılı paragraf eklemek için add_run kullanmak zorundayım, o yüzden önce kalın yazıyı üzerine ekleyeceğim boş bir paragraf açıyorum:
boldparg.add_run('Bu Paragraf Kalın').bold = True


document.save('savetest2.docx')
os.system("start savetest2.docx")